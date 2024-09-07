from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import os

# Ruta del documento de origen
archivo_origen = r'C:\Users\rayoe\OneDrive\Escritorio\MS\Tabla_Personal_Medico.docx'
doc_origen = Document(archivo_origen)

 # Crear un nuevo documento
doc_nuevo = Document()

# Verificar si se cargó correctamente el documento
if not doc_origen.tables:
    print("No se encontraron tablas en el documento.")
else:
    print(f"Se encontró {len(doc_origen.tables)} tabla(s) en el documento.")

# Extraer datos de la tabla del documento original
tabla = doc_origen.tables[0]  # Selecciona la primera tabla

for i, fila in enumerate(tabla.rows[1:], start=1):
    # Suponiendo que los datos relevantes están en la primera fila (ajusta según sea necesario)
    subfijo = fila.cells[0].text
    nombre_completo = fila.cells[1].text + ' ' + fila.cells[2].text + ' ' + fila.cells[3].text
    nombre_archivo = fila.cells[1].text + '_' + fila.cells[2].text + '_' + fila.cells[3].text
    especialidad = fila.cells[4].text
    departamento = fila.cells[5].text

    print(f"Procesando registro {i}: {nombre_completo}, {especialidad}, {departamento}")


    # Crear el encabezado de la carta
    doc_nuevo.add_paragraph('Mazatlán, Sinaloa. A 8 de Julio de 2024', style='Normal').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc_nuevo.add_paragraph()
    doc_nuevo.add_paragraph('Hospital General del Estado de Sinaloa', style='Normal')

    # Añadir el texto del nombre completo
    parrafo_nombre = doc_nuevo.add_paragraph()
    parrafo_nombre.add_run(f'Estimado(a) {subfijo}  ')
    parrafo_nombre.add_run(nombre_completo).bold = True

    # Añadir el texto de la especialidad
    parrafo_especialidad = doc_nuevo.add_paragraph()
    parrafo_especialidad.add_run(f'Con especialidad en  ')
    parrafo_especialidad.add_run(especialidad).bold = True

    # Añadir el texto del departamento
    parrafo_departamento = doc_nuevo.add_paragraph()
    parrafo_departamento.add_run(f'Miembro del departamento de  ')
    parrafo_departamento.add_run(departamento).bold = True

    # Añadir el texto Presente
    parrafo_presente = doc_nuevo.add_paragraph()
    run_presente = parrafo_presente.add_run('PRESENTE')
    run_presente.bold = True
    parrafo_presente.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Añadir el cuerpo de la carta
    carta_texto = (
        "Por medio de la siguiente carta circular, es un placer para la institución invitarlo a usted a "
        "participar en nuestro Primer Congreso Médico de Especialidades, el cual será llevado a cabo los días "
        "28, 29 y 30 de Agosto de 2024 en un horario de las 9:00 a.m. a las 5:00 p.m., en el Hotel DoubleTree By Hilton, "
        "ubicado en Av Camarón Sábalo 905, Zona Dorada, en la ciudad de Mazatlán.\n\n"
        "El objetivo del congreso es mostrar los últimos avances médicos en temas relacionados con la salud, "
        "a través del uso de nuevas tecnologías y avances en la nano medicina.\n\n"
        "Será importante que nos confirme su asistencia al mismo, antes del día 25 de Julio del presente año, para "
        "reservar su habitación de hotel y acceso al evento.\n\n"
        "Confiando en verle en la inauguración, le envío un cordial saludo.\n"
    )
    doc_nuevo.add_paragraph(carta_texto, style='Normal')

    # Añadir la firma
    parrafo_atentamente = doc_nuevo.add_paragraph()
    run_atentamente = parrafo_atentamente.add_run('Atentamente\n')
    run_atentamente.bold = True
    parrafo_atentamente.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Añadir una imagen y centrarla
    ruta_imagen = r'C:\Users\rayoe\OneDrive\Escritorio\MS\firmaJL.jpeg'  # Ruta de la imagen
    parrafo_imagen = doc_nuevo.add_paragraph()
    parrafo_imagen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parrafo_imagen.add_run().add_picture(ruta_imagen, width=Inches(2))  # Ajusta el tamaño de la imagen según sea necesario
    
    parrafo_linea = doc_nuevo.add_paragraph()
    run_linea = parrafo_linea.add_run('_________________________________________________')
    run_linea.bold = True
    parrafo_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER

    parrafo_nombreDirector = doc_nuevo.add_paragraph()
    run_nombreDirector = parrafo_nombreDirector.add_run('Dr. José Luis Meza Herrera.')
    run_nombreDirector.bold = True
    parrafo_nombreDirector.alignment = WD_ALIGN_PARAGRAPH.CENTER

    parrafo_puesto = doc_nuevo.add_paragraph()
    run_puesto = parrafo_puesto.add_run('Director del Hospital General del Estado de Sinaloa.')
    run_puesto.bold = True
    parrafo_puesto.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Añadir un salto de página después de cada carta
    if i < len(tabla.rows) - 1:
        doc_nuevo.add_page_break()

# Guardar el nuevo documento
ruta_guardado = fr'C:\Users\rayoe\OneDrive\Escritorio\MS\Carta_Invitacion.docx'
    
# Verificar si la ruta de guardado existe
if not os.path.exists(os.path.dirname(ruta_guardado)):
        print(f"Directorio de guardado no existe: {os.path.dirname(ruta_guardado)}")
else:
    doc_nuevo.save(ruta_guardado)
    print(f"Documento guardado en: {ruta_guardado}")